from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches


# Function to add a heading with a theme color
def add_heading(doc, text, level, color=None, font_size=12):
    heading = doc.add_heading(text, level=level)
    if color:
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(*color)
    run.font.size = Pt(font_size)


# Function to add content with bold emphasis on key terms
def add_content(doc, theme, content):
    p = doc.add_paragraph()
    p.add_run(f"{theme}: ").bold = True
    for line in content:
        p.add_run(f"\n  - {line}")


# Function to add content in an automation script-like style
def add_automation_script_style(doc, content):
    p = doc.add_paragraph()
    for line in content:
        p.add_run(line + "\n")


# Initialize the document
doc = Document()

# Add a diplomatic and versatile header with a mention of ML/AI
add_heading(doc, "DevSecOps Automation Engineer | CI/CD & Security Automation Specialist | Cloud Infrastructure & AI/ML Enthusiast", level=1, color=(0, 102, 204))

# Stage 1: Dockerized Introduction 🐳
add_heading(doc, "Stage 1: Dockerized Introduction 🐳", level=1, color=(0, 102, 204), font_size=14)
doc.add_paragraph("```\nFROM Akamai-Technologies:Senior-SDET\n")
doc.add_paragraph("MAINTAINER Rajeev_Patil")
doc.add_paragraph("LABEL Role=\"Senior SDET | DevSecOps Expert | AI/ML Enthusiast\"")
doc.add_paragraph("LABEL Certifications=\"CEH, AZ-500\"")
doc.add_paragraph("ENV Expertise=\"CI/CD Automation | Cloud Security | AI/ML\"")
doc.add_paragraph("ENV CurrentRole=\"Designing automation frameworks & cloud security protocols at Akamai Technologies\"")
doc.add_paragraph("ENV Tools=\"Kubernetes, Docker\"")
doc.add_paragraph("\n# Continuous learning and driving innovation in DevSecOps & AI/ML.\nRUN echo \"Passionate about securing and optimizing cloud platforms with cutting-edge technologies.\"\n```\n")

doc.add_paragraph(
    "With over 11 years of experience in software testing, automation, and DevSecOps, I specialize in integrating security into CI/CD pipelines, optimizing cloud infrastructure, and leveraging AI/ML for system performance. Currently at Akamai Technologies, I design advanced automation frameworks, cloud security protocols, and scalable solutions using Kubernetes and Docker. Holding certifications like CEH and AZ-500, I am committed to continuous learning and driving innovation in DevSecOps and AI/ML."
)



# Stage 2: Skills (Kubernetes Themed) ☸️
add_heading(doc, "Stage 2: skills-deployment.yaml ☸️", level=1, color=(0, 153, 51), font_size=14)
doc.add_paragraph("```yaml\napiVersion: v1\nkind: SkillSet\nmetadata:\n  name: Skills")
doc.add_paragraph("spec:\n  - Programming: Python, Shell Scripting, SQL, Linux, REST API")
doc.add_paragraph("  - Test Automation: Pytest, Robot Framework, Selenium, Locust")
doc.add_paragraph(
    "  - Cloud & Security: Azure, Vulnerability Scanning, Container Security, WIZ, Trivy, Codeclimate")
doc.add_paragraph("  - containers_and_orchestration : Docker, Kubernetes")
doc.add_paragraph("  - AI/ML: TensorFlow, Data Science, Pandas, NumPy")
doc.add_paragraph("  - devops_tools: YAML, GitLab CI/CD, Jenkins, Argo CD basics \n```\n")

# Stage 3: Work History (Automation Script) 🐍
add_heading(doc, "Stage 3: Work History (Automation Script) 🐍", level=1, color=(255, 102, 0), font_size=14)

work_history_script = [
    "### Starting Work History Automation Logs",
    "> [INFO] Loading work history...",
    "# Company: Akamai Technologies - Armada-NG",
    "  * Role: Senior SDET, DevSecOps Automation",
    "  * Technologies Used: Python, GitLab APIs, Robot Framework, Docker, Kubernetes",
    "> [INFO] Initiating work for Akamai Technologies...",
    "  - Developed and optimized security-first automation framework.",
    "  - Integrated DevSecOps practices into CI/CD pipelines.",
    "  - Automated LOKI provisioning on Kubernetes.",
    "  - Utilized AI/ML modules for actionable insights into build quality.",
    "> [SUCCESS] Akamai Technologies work history processed.",
    "",
    "# Company: Safran Engineering",
    "  * Role: Software Development Engineer in Test",
    "> [INFO] Initiating work for Safran Engineering...",
    "  - Automated 200+ test scenarios for In-flight Entertainment features.",
    "  - Enhanced testing efficiency for backend features with Python and Flask.",
    "  - Conducted data-driven analysis with Pandas.",
    "  - Developed regression and performance testing frameworks.",
    "> [SUCCESS] Safran Engineering work history processed.",
    "",
    "# Company: Infoblox Pvt Ltd",
    "  * Role: Software Engineer II",
    "> [INFO] Initiating work for Infoblox Pvt Ltd...",
    "  - Security testing using Nessus and Burp Suite.",
    "  - API Automation with Python.",
    "  - Enhanced testing for microservices-based cloud platforms.",
    "  - Integrated validation tools and CI/CD pipelines.",
    "> [SUCCESS] Infoblox Pvt Ltd work history processed.",
    "",
    "# Company: CenturyLink Technologies",
    "  * Role: Full Stack QA–UI/Backend on Paypal Integration",
    "> [INFO] Initiating work for CenturyLink Technologies...",
    "  - Developed testing frameworks for API and UI validation.",
    "  - Collaborated with cross-functional teams for deployment strategies.",
    "  - Performed testing for mobile platforms on iOS and Android.",
    "> [SUCCESS] CenturyLink Technologies work history processed.",
    "",
    "# Company: Thomson Reuters",
    "  * Role: Systems Engineer (Manual Testing/Support)",
    "> [INFO] Initiating work for Thomson Reuters...",
    "  - Created test plans and case scenarios.",
    "  - Utilized Jira for issue tracking and team reporting.",
    "  - Supported the QA team and ensured Agile methodology adherence.",
    "> [SUCCESS] Thomson Reuters work history processed.",
    "### Work History Automation Logs Complete"
]

add_automation_script_style(doc, work_history_script)

# Stage 4: Certifications (Security Scanning) 🛡️
add_heading(doc, "Stage 4: Certifications (Security Scanning) 🛡️", level=1, color=(204, 0, 0), font_size=14)
doc.add_paragraph("```\nScanning for certifications:\n")
certifications = [
    "Certified Ethical Hacker (CEH)",
    "Microsoft Certified: Azure Security Engineer Associate (AZ-500)",
    "Microsoft Cybersecurity Architect (SC-100) - In Progress",
    "Python Certification – Udemy"
]
add_content(doc, "Certifications", certifications)
doc.add_paragraph("\nScan Complete! No vulnerabilities found.\n```")


# Stage 5: Demos (GitLab CI/CD) 🦊
add_heading(doc, "Stage 5: Demos (GitLab CI/CD) 🦊", level=1, color=(128, 0, 128), font_size=14)
doc.add_paragraph("```yaml\nstages:\n  - Introduction")
doc.add_paragraph("  - Skills\n  - Work_History\n  - Certifications\n  - Demos")
doc.add_paragraph("\nDemos:\n  stage: demos")
doc.add_paragraph(
    "  script:\n    - echo 'This resume is created using Python Script to demonstrate automation skills , please visit abc.com for more details'")
doc.add_paragraph("  script:\n    - echo 'Upcoming Demonstrations on Security and GitLab CI/CD Practices'")
doc.add_paragraph("    - echo 'Stay tuned for more updates!'\n```\n")

# Stage 6: Academic & Professional Development 🎓
add_heading(doc, "Stage 6: Academic & Professional Development 🎓", level=1, color=(0, 102, 204), font_size=14)

# Education content
education_content = [
    "PG Diploma in Machine Learning & AI – **IIIT Bangalore**",
    "Bachelor of Engineering (B.E.) in Electronics & Communication – **XYZ University**"
]

add_content(doc, "Education", education_content)

# Leadership qualities section
doc.add_paragraph(
    "As a DevSecOps Automation Engineer, I specialize in automating security-first workflows, enhancing CI/CD pipelines, and optimizing cloud infrastructure. I work on integrating security practices into automation processes and ensure robust, scalable solutions across cloud platforms.")
doc.add_paragraph("Explore my complete portfolio and experience on my website, hosted on GitLab.")
doc.add_paragraph("Want to connect and discuss innovative solutions? Visit my website for more details:")
doc.add_paragraph("Personal Website: (https://rajeevpatil24.github.io/rajeevpatil.github.io/)").italic = True
doc.add_paragraph("LinkedIn: [Rajeev Patil](https://www.linkedin.com/in/rajeevpatil-v)").italic = True
doc.add_paragraph("GitHub: [Rajeev Patil GitHub](https://github.com/rajeevpatil24)").italic = True


# Save the document
doc.save("Rajeev_Patil_DevSecOps_Bangalore_Nov2024_test.docx")
print("The CV has been generated as 'Rajeev_Patil_DevSecOps_Nov2024.docx'.")
